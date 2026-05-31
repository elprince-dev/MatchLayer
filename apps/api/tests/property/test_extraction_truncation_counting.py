"""Feature: phase-1-matching — Property 15.

Property 15: Extracted text is truncated and counted consistently.

    *For any* extracted text, the stored ``extracted_text`` length is at most
    ``MATCHLAYER_RESUME_MAX_EXTRACTED_CHARS``, and ``extraction_char_count``
    equals the length of the stored text.

**Validates: Requirements 3.3**

This is the universal companion to the concrete failure-category coverage in
``tests/unit`` (task 7.6). Where those tests pin down the fail-soft vocabulary
(``extraction_timeout`` / ``corrupt_document`` / ``empty_text``), this module
asserts the *success-path* invariant of
:func:`matchlayer_api.services.extraction.extract` across a wide, generated
input space using Hypothesis (>=100 examples).

The property in code form: whenever ``extract`` returns
``status == "succeeded"``,

    len(outcome.text) <= max_extracted_chars      # truncation
    outcome.char_count == len(outcome.text)       # counting

To drive ``extract`` with *real* inputs (no mocks), each example builds an
in-memory ``.docx`` with ``python-docx`` from Hypothesis-generated paragraph
text of varying lengths, then varies ``max_extracted_chars`` — including caps
far smaller than the document so the truncation branch actually fires. DOCX is
the chosen generator because ``python-docx`` can both *write* the fixture and
is what the extractor *reads* (``services/extraction.py`` ``_extract_docx_sync``
iterates ``document.paragraphs``); a pure-Python PDF whose text ``pypdf`` can
re-extract is impractical without an extra writer dependency, and the design
notes the truncation/char-cap logic is shared across both kinds (the caller
applies the exact ``raw_text[:max_extracted_chars]`` cap regardless of kind),
so DOCX coverage exercises the property end to end.

``extract`` is ``async``; following the suite's established pattern
(``tests/property/test_rate_limit_window.py``) the coroutine is driven inside
an :class:`asyncio.Runner` per example rather than relying on an outer event
loop. A generous ``timeout_seconds`` is passed so the wall-clock/timeout branch
never interferes with the truncation invariant under test.
"""

from __future__ import annotations

import asyncio
import string
from collections.abc import Awaitable, Callable
from io import BytesIO

import docx
from hypothesis import given, settings
from hypothesis import strategies as st

from matchlayer_api.services.extraction import ExtractionOutcome, extract

# A generous wall-clock ceiling: large enough that building/parsing a small
# in-memory DOCX never approaches it, so the ``extraction_timeout`` branch
# cannot fire and confound the truncation invariant under test.
_GENEROUS_TIMEOUT_SECONDS = 30.0

# Characters safe to place inside an OOXML (XML 1.0) text run. Every codepoint
# is >= 0x20, so ``lxml`` (under ``python-docx``) serializes them without
# raising on a forbidden control character. A handful of non-ASCII letters are
# included so the property also exercises multi-byte-but-single-codepoint text:
# Python ``len`` counts codepoints, the extractor truncates by codepoint, and
# ``char_count == len(text)`` must hold for those too.
_SAFE_CHARS = string.ascii_letters + string.digits + " .,-_/+#" + "éñüçßΩ你日本語"

# Paragraph text. ``min_size=1`` keeps a paragraph from being empty; the
# letter-heavy alphabet means the joined document is overwhelmingly
# non-whitespace, so most examples reach the ``succeeded`` branch (the only
# branch the property constrains). Whitespace-only draws simply fall through to
# the fail-soft ``empty_text`` outcome and are skipped by the guard below.
_paragraph_text = st.text(alphabet=_SAFE_CHARS, min_size=1, max_size=400)

# A multi-paragraph document. Up to 20 paragraphs of up to 400 chars yields
# documents up to ~8000 chars, comfortably larger than the small end of the
# ``max_extracted_chars`` range below, guaranteeing the truncation branch is
# exercised across the run.
_paragraphs = st.lists(_paragraph_text, min_size=1, max_size=20)

# The character cap to truncate to. The low end (1) forces aggressive
# truncation; the high end can exceed the whole document so the no-truncation
# path is covered too.
_max_extracted_chars = st.integers(min_value=1, max_value=4000)


def _build_docx(paragraphs: list[str]) -> bytes:
    """Serialize ``paragraphs`` into a real in-memory ``.docx`` byte payload.

    Uses ``python-docx`` so the fixture is structurally identical to what the
    extractor reads in production — no hand-rolled OOXML, no mocks.
    """
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _run_sync(coro_factory: Callable[[], Awaitable[None]]) -> None:
    """Drive an async test body via :class:`asyncio.Runner`.

    Mirrors ``tests/property/test_rate_limit_window.py``: ``Runner`` closes the
    event loop deterministically on ``__exit__``, avoiding the
    ``ResourceWarning("unclosed event loop")`` that bare ``asyncio.run`` can
    leak in teardown (the suite runs under ``filterwarnings = ["error"]``).
    """
    with asyncio.Runner() as runner:
        runner.run(coro_factory())


@settings(max_examples=150, deadline=None)
@given(paragraphs=_paragraphs, max_extracted_chars=_max_extracted_chars)
def test_succeeded_extraction_is_truncated_and_counted(
    paragraphs: list[str], max_extracted_chars: int
) -> None:
    """A succeeded extraction never exceeds the cap and counts itself exactly.

    Property 15 over a generated space of real DOCX documents and character
    caps. When ``extract`` reports ``succeeded``, the retained text length is
    at most ``max_extracted_chars`` and ``char_count`` equals that length. A
    failed outcome (e.g. a whitespace-only draw → ``empty_text``) carries no
    text and is outside the property's scope, so it is skipped.
    """
    data = _build_docx(paragraphs)

    async def _run() -> None:
        outcome: ExtractionOutcome = await extract(
            data,
            "docx",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=max_extracted_chars,
        )

        if outcome.status != "succeeded":
            # Fail-soft outcomes (text is None) are not what Property 15
            # constrains; Property 16 (task 7.5's sibling) owns those.
            assert outcome.text is None
            assert outcome.char_count is None
            return

        assert outcome.text is not None
        assert outcome.char_count is not None
        # Truncation: never longer than the configured cap.
        assert len(outcome.text) <= max_extracted_chars
        # Counting: the recorded count is exactly the stored text length.
        assert outcome.char_count == len(outcome.text)

    _run_sync(_run)


@settings(max_examples=100, deadline=None)
@given(
    body_length=st.integers(min_value=200, max_value=5000),
    max_extracted_chars=st.integers(min_value=1, max_value=150),
)
def test_truncation_is_tight_when_document_exceeds_cap(
    body_length: int, max_extracted_chars: int
) -> None:
    """When the document is larger than the cap, the text is cut to exactly it.

    A single long, all-non-whitespace paragraph makes the extractor's behaviour
    deterministic: the joined text is that one paragraph (no inter-paragraph
    newlines), it is strictly longer than the cap, and it strips to non-empty —
    so the outcome must be ``succeeded`` with ``len(text) == max_extracted_chars``
    and ``char_count == max_extracted_chars``. This pins the truncation-triggered
    edge that the broader property only bounds with ``<=``.
    """
    # ``body_length`` always exceeds ``max_extracted_chars`` (200 > 150), so
    # truncation is guaranteed to fire.
    paragraph = "a" * body_length
    data = _build_docx([paragraph])

    async def _run() -> None:
        outcome: ExtractionOutcome = await extract(
            data,
            "docx",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=max_extracted_chars,
        )

        assert outcome.status == "succeeded"
        assert outcome.text is not None
        assert outcome.char_count is not None
        assert len(outcome.text) == max_extracted_chars
        assert outcome.char_count == max_extracted_chars

    _run_sync(_run)
