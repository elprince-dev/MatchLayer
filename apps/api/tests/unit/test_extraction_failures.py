"""Unit tests for the ``Resume_Extractor`` fail-soft failure categories.

Task 7.6. Validates Requirements 3.5, 3.7.

Concrete-example coverage for the three failure verdicts
:func:`matchlayer_api.services.extraction.extract` can return, pinning down the
fail-soft vocabulary the property suite only bounds in aggregate
(``tests/property/test_extraction_truncation_counting.py`` owns the
success-path truncation/counting invariant, Property 15):

* ``extraction_timeout`` — the cooperative wall-clock deadline trips during
  page/paragraph iteration, or the outer ``asyncio.wait_for`` guard fires
  (Requirement 3.2).
* ``corrupt_document`` — ``pypdf`` / ``python-docx`` raise a parser error on
  bytes that are not a valid PDF/DOCX of the claimed ``kind`` (Requirement
  3.5, fail-soft: this must NOT raise into the request path).
* ``empty_text`` — a structurally valid document yields only whitespace
  (Requirement 3.5).

**Where the PII-safe failure log actually lives (Requirements 3.6, 3.7).**
:func:`extract` deliberately does **no logging at all**. Per the module
docstring of ``services/extraction.py``, the function has no Resume ``id`` (the
``resumes`` row is created by the ``Resume_Service`` *after* storage), so the
single structured "extraction failed" event that names the ``failure_category``
and the Resume ``id`` — and never the file bytes or extracted text — is emitted
by the ``Resume_Service`` from the returned :class:`ExtractionOutcome`. These
tests therefore scope the Requirement 3.7 assertion to what *this* layer is
responsible for: ``extract`` returns the right ``failure_category`` for each
case and emits **no** structured (structlog) log line — in particular none
carrying the input bytes or extracted text. The ``Resume_Service``-side
structured log (category + Resume id, never bytes/text) is exercised by the
"never logged" PII negatives in task 10.7's integration tests.

These tests use **real** inputs — no mocks: garbage byte payloads for the
parser-error path and ``python-docx``-authored documents for the
whitespace-only and timeout paths — so the assertions exercise the production
code path end to end.

References:
* Requirement 3.5 (fail-soft: failure never returns a 5xx / never raises).
* Requirement 3.7 (failure logging names the category + Resume id only, never
  bytes or text — owned by the ``Resume_Service``; see task 10.7).
* Design §"Resume_Extractor".
"""

from __future__ import annotations

from collections.abc import MutableMapping
from io import BytesIO
from typing import Any

import docx
import structlog

from matchlayer_api.services.extraction import ExtractionOutcome, extract

# Bounds passed explicitly (the function reads neither from settings); a
# generous char cap keeps truncation out of these failure-path tests.
_MAX_EXTRACTED_CHARS = 200_000
# A generous wall-clock ceiling for the non-timeout cases, large enough that
# parsing a tiny in-memory document never approaches it.
_GENEROUS_TIMEOUT_SECONDS = 30.0
# A zero deadline forces the timeout path: both the cooperative in-thread
# ``monotonic() > deadline`` check and the outer ``asyncio.wait_for(timeout=0)``
# guard map to ``extraction_timeout``.
_ZERO_TIMEOUT_SECONDS = 0.0


def _build_docx(paragraphs: list[str]) -> bytes:
    """Serialize ``paragraphs`` into a real in-memory ``.docx`` byte payload.

    Uses ``python-docx`` so the fixture is structurally identical to what the
    extractor reads in production (``_extract_docx_sync`` iterates
    ``document.paragraphs``) — no hand-rolled OOXML, no mocks.
    """
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _assert_no_structlog_pii(
    captured: list[MutableMapping[str, Any]], *, markers: list[str]
) -> None:
    """Assert ``extract`` emitted no structured log carrying bytes/text.

    ``extract`` is expected to emit *nothing* (the ``Resume_Service`` owns the
    failure log). The primary assertion is therefore that ``captured`` is
    empty. The substring scan is defence in depth: should the logging owner
    ever move into this layer, a regression that echoes a recognizable slice of
    the input bytes or extracted text into a log event would still trip here.
    """
    # Primary contract: this layer does not log at all (Requirement 3.7 places
    # the failure log on the Resume_Service; see task 10.7).
    assert captured == [], f"extract() should emit no structured logs, got: {captured!r}"

    # Defence in depth: no captured event, serialized, may contain any marker
    # planted in the input bytes or document text.
    haystack = repr(captured)
    for marker in markers:
        assert marker not in haystack, f"log output leaked input content: {marker!r}"


async def test_corrupt_pdf_bytes_fail_soft_as_corrupt_document() -> None:
    """Garbage bytes with ``kind='pdf'`` fail soft as ``corrupt_document``.

    ``pypdf`` rejects the invalid header / missing EOF, and the fail-soft
    contract (Requirement 3.5) converts that parser error into a failed
    outcome rather than letting it raise into the request path. Reaching the
    assertions below *is* the no-raise assertion: any propagated exception
    would surface as a test error.

    Validates: Requirements 3.5, 3.7.
    """
    # A recognizable ASCII marker is embedded so the PII-safety scan is
    # meaningful: if any log echoed these "file bytes", it would be detected.
    marker = "CORRUPTPDFRESUMEBYTES"
    garbage = (marker + " not a real pdf \x00\x01\x02\x03 ").encode("latin-1") * 16

    with structlog.testing.capture_logs() as captured:
        outcome: ExtractionOutcome = await extract(
            garbage,
            "pdf",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=_MAX_EXTRACTED_CHARS,
        )

    assert outcome.status == "failed"
    assert outcome.failure_category == "corrupt_document"
    # Fail-soft: a failed outcome carries no text or count (Requirement 3.5).
    assert outcome.text is None
    assert outcome.char_count is None
    _assert_no_structlog_pii(captured, markers=[marker])


async def test_corrupt_docx_bytes_fail_soft_as_corrupt_document() -> None:
    """Non-archive bytes with ``kind='docx'`` fail soft as ``corrupt_document``.

    A payload that is not a readable OOXML ZIP makes ``python-docx`` raise; the
    fail-soft contract maps it to ``corrupt_document`` without raising.

    Validates: Requirements 3.5, 3.7.
    """
    marker = "CORRUPTDOCXRESUMEBYTES"
    garbage = (marker + " definitely not a zip ").encode("latin-1") * 16

    with structlog.testing.capture_logs() as captured:
        outcome: ExtractionOutcome = await extract(
            garbage,
            "docx",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=_MAX_EXTRACTED_CHARS,
        )

    assert outcome.status == "failed"
    assert outcome.failure_category == "corrupt_document"
    assert outcome.text is None
    assert outcome.char_count is None
    _assert_no_structlog_pii(captured, markers=[marker])


async def test_whitespace_only_docx_fails_as_empty_text() -> None:
    """A structurally valid DOCX that holds only whitespace → ``empty_text``.

    The document parses cleanly (no ``corrupt_document``) but the extracted
    text strips to nothing, which the fail-soft contract records as
    ``empty_text`` (Requirement 3.5) — the resume is simply not matchable.

    Validates: Requirements 3.5, 3.7.
    """
    # A valid .docx whose paragraphs contain only spaces/tabs/newlines.
    empty_docx = _build_docx(["   ", "\t", "  \u00a0 ", ""])

    with structlog.testing.capture_logs() as captured:
        outcome: ExtractionOutcome = await extract(
            empty_docx,
            "docx",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=_MAX_EXTRACTED_CHARS,
        )

    assert outcome.status == "failed"
    assert outcome.failure_category == "empty_text"
    assert outcome.text is None
    assert outcome.char_count is None
    # No meaningful text to leak, but assert the no-log contract holds.
    _assert_no_structlog_pii(captured, markers=[])


async def test_zero_deadline_fails_as_extraction_timeout() -> None:
    """A zero wall-clock budget over a real DOCX → ``extraction_timeout``.

    With ``timeout_seconds=0`` the precomputed monotonic deadline has already
    passed by the time the worker thread begins iterating paragraphs, so the
    cooperative deadline check aborts (and/or the outer ``asyncio.wait_for``
    guard fires) — both map to ``extraction_timeout`` (Requirement 3.2). A
    multi-paragraph document guarantees the iteration loop runs at least one
    deadline check.

    Validates: Requirements 3.5, 3.7.
    """
    marker = "CONFIDENTIALTIMEOUTPARAGRAPH"
    paragraphs = [f"{marker} line {i} with several words to extract" for i in range(64)]
    big_docx = _build_docx(paragraphs)

    with structlog.testing.capture_logs() as captured:
        outcome: ExtractionOutcome = await extract(
            big_docx,
            "docx",
            timeout_seconds=_ZERO_TIMEOUT_SECONDS,
            max_extracted_chars=_MAX_EXTRACTED_CHARS,
        )

    assert outcome.status == "failed"
    assert outcome.failure_category == "extraction_timeout"
    assert outcome.text is None
    assert outcome.char_count is None
    # The document text (potential PII) must never reach a log line.
    _assert_no_structlog_pii(captured, markers=[marker])


async def test_extract_emits_no_structured_log_on_failure() -> None:
    """``extract`` itself logs nothing; the Resume_Service owns the failure log.

    Consolidates the Requirement 3.7 ownership boundary into one explicit
    assertion: across a representative failure (corrupt PDF), :func:`extract`
    produces a categorized outcome but emits **zero** structlog events. The
    structured "extraction failed" line that names the ``failure_category`` and
    the Resume ``id`` (and never the bytes or text) is emitted by the
    ``Resume_Service`` and is covered by task 10.7's integration tests.

    Validates: Requirements 3.6, 3.7.
    """
    garbage = b"%PDF-1.4 truncated, no xref, no trailer, no eof marker"

    with structlog.testing.capture_logs() as captured:
        outcome: ExtractionOutcome = await extract(
            garbage,
            "pdf",
            timeout_seconds=_GENEROUS_TIMEOUT_SECONDS,
            max_extracted_chars=_MAX_EXTRACTED_CHARS,
        )

    assert outcome.status == "failed"
    assert outcome.failure_category is not None
    assert captured == [], (
        "extract() must not emit structured logs; the failure log naming "
        "failure_category + resume id is the Resume_Service's responsibility "
        "(Requirement 3.7, task 10.7)."
    )
