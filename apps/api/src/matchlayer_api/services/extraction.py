"""Bounded, fail-soft resume text extraction (the ``Resume_Extractor``).

Converts an uploaded PDF or DOCX into plain UTF-8 text under hard resource
bounds, then hands the result back to the ``Resume_Service`` as a small
:class:`ExtractionOutcome` value. Extraction is a one-way transformation, not
a reversible parser — there is no round-trip or pretty-printer contract here
(see the ``Resume_Extractor`` glossary entry).

Two distinct safety contracts live in this module, and they deliberately
behave differently:

* :func:`extract` is **fail-soft** (Requirement 3.5). A corrupt, empty,
  whitespace-only, or slow document never raises into the request path; it
  yields ``ExtractionOutcome(status="failed", text=None, ...)`` with a
  ``failure_category`` so the ``Resume_Service`` can persist the resume with
  ``extraction_status='failed'`` and still return its ``201`` — the upload
  succeeded even though the text could not be read. The resume is simply not
  matchable until re-uploaded.
* :func:`guard_docx_archive` is **fail-fast** (Requirement 2.4). It is the
  decompression-bomb defense from the ``security.md`` file-upload threat
  model and runs *before* storage and extraction: a DOCX whose declared
  uncompressed size or entry count is over the configured ceilings is
  rejected with :class:`~matchlayer_api.core.errors.MalformedUploadError`
  (HTTP 422 ``malformed_upload``) so no object is written and no row is
  persisted. This is intentionally not a fail-soft ``failure_category`` —
  the requirement is to *refuse* the upload, not to store it as unparseable.

Resource bounds (Requirements 3.1-3.3):

* **Wall-clock timeout.** The blocking ``pypdf`` / ``python-docx`` work runs
  in a worker thread via :func:`fastapi.concurrency.run_in_threadpool` under
  an :func:`asyncio.wait_for` guard. Because a Python thread cannot be force-
  killed (and ``run_in_threadpool`` does not abandon it on cancellation),
  the *effective* bound is a **cooperative wall-clock check**: the sync
  extractor compares :func:`time.monotonic` against a precomputed deadline
  after each page/paragraph and aborts early once the deadline passes. The
  ``asyncio.wait_for`` is the outer asyncio-level guard; the cooperative
  check is what actually stops the thread from burning CPU past the
  deadline. Sandboxed/queued parsing is deferred to Phase 4 per
  ``security.md``.
* **Character cap.** Extracted text is truncated to at most
  ``max_extracted_chars`` characters and the retained count is reported, so
  a pathological document cannot produce an unbounded string.

PII discipline (Requirements 3.6, 3.7): this module **never logs** the file
bytes or any extracted text — both are Restricted PII per ``security.md``.
It does not log at all: it has no Resume ``id`` (the row is created by the
``Resume_Service`` *after* storage), so the structured "extraction failed"
event that names the ``failure_category`` and the resume ``id`` is emitted by
the ``Resume_Service`` from the returned :class:`ExtractionOutcome`. Keeping
logging in the caller avoids a second, id-less log line here and keeps the
"resume id and failure_category only" guarantee in one place.

Design reference: §"Resume_Extractor".
Requirements covered: 2.4, 3.1, 3.2, 3.3, 3.4, 3.5, 3.7.
"""

from __future__ import annotations

import asyncio
import zipfile
from dataclasses import dataclass
from io import BytesIO
from time import monotonic
from typing import Literal

# ``python-docx`` imports as ``docx``. Recent releases ship a ``py.typed``
# marker, so mypy resolves its types; we still coerce paragraph ``.text`` with
# ``str`` below to keep a concrete ``str`` return under ``warn_return_any``.
import docx
from fastapi.concurrency import run_in_threadpool
from pypdf import PdfReader

from matchlayer_api.core.errors import MalformedUploadError

__all__ = [
    "ExtractionOutcome",
    "FailureCategory",
    "ResumeKind",
    "extract",
    "guard_docx_archive",
]

# The two file kinds the upload surface accepts. Mirrors the Mime_Validator
# verdict (``core/mime.py``) and ``Resume_Storage.ResumeKind``
# (``core/storage.py``); defined locally so this module stays decoupled from
# the storage/boto3 import surface — extraction reads bytes, it never touches
# object storage.
ResumeKind = Literal["pdf", "docx"]

# The fail-soft failure vocabulary recorded on a failed ``ExtractionOutcome``
# (Requirement 3.7). The DOCX zip-bomb case is deliberately *not* here: it is
# a fail-fast ``malformed_upload`` (HTTP 422) raised by
# :func:`guard_docx_archive` before persistence, not a stored failure state.
FailureCategory = Literal["extraction_timeout", "corrupt_document", "empty_text"]


@dataclass(frozen=True, slots=True)
class ExtractionOutcome:
    """The result of attempting to extract text from a resume file.

    A discriminated value the ``Resume_Service`` maps directly onto the
    ``resumes`` extraction columns:

    * On success: ``status="succeeded"``, ``text`` is the truncated UTF-8
      text (guaranteed to contain at least one non-whitespace character),
      ``char_count`` equals ``len(text)``, and ``failure_category`` is
      ``None``.
    * On failure: ``status="failed"``, ``text`` is ``None``, ``char_count``
      is ``None``, and ``failure_category`` names why (Requirement 3.5).
    """

    status: Literal["succeeded", "failed"]
    text: str | None
    char_count: int | None
    failure_category: FailureCategory | None


class _ExtractionTimeoutError(Exception):
    """Internal signal: the cooperative wall-clock deadline was reached.

    Raised inside the worker thread between pages/paragraphs so the thread
    stops promptly instead of running until the document is exhausted. It is
    caught in :func:`extract` and mapped to ``failure_category`` =
    ``"extraction_timeout"`` — it never escapes this module.
    """


def _failed(category: FailureCategory) -> ExtractionOutcome:
    """Build the canonical fail-soft outcome for ``category``."""
    return ExtractionOutcome(
        status="failed",
        text=None,
        char_count=None,
        failure_category=category,
    )


def _extract_pdf_sync(data: bytes, *, deadline: float, max_chars: int) -> str:
    """Extract PDF text page-by-page under a cooperative deadline and char cap.

    Runs in a worker thread. Raises :class:`_ExtractionTimeoutError` if the
    monotonic deadline passes between pages, and stops early once the
    accumulated length reaches ``max_chars`` (the caller truncates to the
    exact cap). Any ``pypdf`` parsing error propagates and is mapped to
    ``corrupt_document`` by the caller.

    Args:
        data: Raw PDF bytes (Restricted PII — never logged).
        deadline: A :func:`time.monotonic` timestamp; iteration aborts once
            ``monotonic()`` exceeds it.
        max_chars: Stop accumulating once this many characters are gathered.

    Returns:
        The concatenated page text (may exceed ``max_chars`` by at most the
        final page's contribution; the caller applies the exact cap).
    """
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        if monotonic() > deadline:
            raise _ExtractionTimeoutError
        page_text = page.extract_text() or ""
        parts.append(page_text)
        total += len(page_text)
        if total >= max_chars:
            break
    return "\n".join(parts)


def _extract_docx_sync(data: bytes, *, deadline: float, max_chars: int) -> str:
    """Extract DOCX paragraph text under a cooperative deadline and char cap.

    Runs in a worker thread; mirrors :func:`_extract_pdf_sync` but iterates
    ``document.paragraphs``. Each paragraph's ``.text`` is coerced with
    :class:`str` to keep a typed ``str`` return under ``warn_return_any``.

    Args:
        data: Raw DOCX bytes (Restricted PII — never logged). The caller is
            expected to have already run :func:`guard_docx_archive` against
            these bytes (Requirement 2.4).
        deadline: A :func:`time.monotonic` timestamp; iteration aborts once
            ``monotonic()`` exceeds it.
        max_chars: Stop accumulating once this many characters are gathered.

    Returns:
        The newline-joined paragraph text.
    """
    document = docx.Document(BytesIO(data))
    parts: list[str] = []
    total = 0
    for paragraph in document.paragraphs:
        if monotonic() > deadline:
            raise _ExtractionTimeoutError
        paragraph_text = str(paragraph.text)
        parts.append(paragraph_text)
        total += len(paragraph_text)
        if total >= max_chars:
            break
    return "\n".join(parts)


async def extract(
    data: bytes,
    kind: ResumeKind,
    *,
    timeout_seconds: float,
    max_extracted_chars: int,
) -> ExtractionOutcome:
    """Extract plain UTF-8 text from a resume file, fail-soft and bounded.

    Dispatches to the PDF or DOCX sync extractor in a worker thread under an
    :func:`asyncio.wait_for` wall-clock guard, with a cooperative
    :func:`time.monotonic` deadline checked inside the thread (Requirements
    3.1, 3.2). The result is truncated to ``max_extracted_chars`` and the
    retained count is reported (Requirement 3.3).

    This function **never raises** because extraction failed (Requirement
    3.5): a timeout, a parser error, or empty/whitespace-only text all return
    a failed :class:`ExtractionOutcome`. (It will still propagate
    :class:`asyncio.CancelledError` if the surrounding request is cancelled —
    that is task cancellation, not an extraction failure.)

    The bounds are accepted as parameters rather than read from settings here
    so the function stays a pure, directly-testable unit; the
    ``Resume_Service`` passes the configured
    ``MATCHLAYER_RESUME_EXTRACTION_TIMEOUT_SECONDS`` and
    ``MATCHLAYER_RESUME_MAX_EXTRACTED_CHARS`` values in.

    Args:
        data: The raw file bytes. Restricted PII — never logged or echoed.
        kind: ``"pdf"`` or ``"docx"`` — the Mime_Validator's magic-byte
            verdict, never the client ``Content-Type`` or extension.
        timeout_seconds: Wall-clock ceiling for the whole extraction.
        max_extracted_chars: Maximum number of characters to retain.

    Returns:
        ``ExtractionOutcome(status="succeeded", ...)`` with truncated text and
        a matching ``char_count`` when at least one non-whitespace character
        was extracted; otherwise a failed outcome whose ``failure_category``
        is ``"extraction_timeout"``, ``"corrupt_document"``, or
        ``"empty_text"``.
    """
    deadline = monotonic() + timeout_seconds
    if kind == "pdf":
        coro = run_in_threadpool(
            _extract_pdf_sync, data, deadline=deadline, max_chars=max_extracted_chars
        )
    else:
        coro = run_in_threadpool(
            _extract_docx_sync, data, deadline=deadline, max_chars=max_extracted_chars
        )

    try:
        raw_text = await asyncio.wait_for(coro, timeout=timeout_seconds)
    except (TimeoutError, _ExtractionTimeoutError):
        # Wall-clock bound hit: the asyncio guard fired, or the cooperative
        # in-thread deadline check aborted iteration (Requirement 3.2).
        return _failed("extraction_timeout")
    except Exception:
        # Fail-soft contract (Requirement 3.5): any pypdf / python-docx
        # parsing error (corrupt header, broken xref, bad zip, encrypted
        # document, ...) becomes a failed outcome, never a 5xx. The bytes are
        # PII and are deliberately not included in the (caller-emitted) log.
        # ``CancelledError`` is a BaseException and is intentionally NOT
        # swallowed by this ``Exception`` clause.
        return _failed("corrupt_document")

    # Apply the exact character cap (the sync extractor may overrun by the
    # final page/paragraph). char_count == len(text) <= max_extracted_chars.
    text = raw_text[:max_extracted_chars]
    if not text.strip():
        # Empty or whitespace-only extraction is a failure (Requirement 3.5):
        # a scanned/image-only PDF or an empty document yields no usable text.
        return _failed("empty_text")

    return ExtractionOutcome(
        status="succeeded",
        text=text,
        char_count=len(text),
        failure_category=None,
    )


def guard_docx_archive(
    data: bytes,
    *,
    max_decompressed_bytes: int,
    max_archive_entries: int,
) -> None:
    """Reject a decompression-bomb DOCX before storage or extraction.

    A DOCX is an OOXML ZIP container. The stdlib :mod:`zipfile` central
    directory exposes each entry's declared uncompressed size
    (:attr:`zipfile.ZipInfo.file_size`) and the full entry list *without*
    decompressing anything, so this guard runs cheaply before any object is
    written or any text is extracted (Requirement 2.4, ``security.md``
    file-upload threat model).

    Unlike :func:`extract`, this guard is **fail-fast**: on a violation it
    raises :class:`~matchlayer_api.core.errors.MalformedUploadError` (HTTP 422
    ``malformed_upload``) so the ``Resume_Service`` refuses the upload and
    persists nothing. The bounds are passed in by the service from
    ``MATCHLAYER_RESUME_MAX_DECOMPRESSED_BYTES`` and
    ``MATCHLAYER_RESUME_MAX_ARCHIVE_ENTRIES``.

    Args:
        data: The raw DOCX bytes, already confirmed as a DOCX by the
            Mime_Validator. Restricted PII — never logged.
        max_decompressed_bytes: Ceiling on the sum of all entries' declared
            uncompressed sizes.
        max_archive_entries: Ceiling on the number of archive entries.

    Raises:
        MalformedUploadError: If the archive cannot be opened, has more than
            ``max_archive_entries`` entries, or its total declared
            uncompressed size exceeds ``max_decompressed_bytes``. The detail
            string names only the limit — never any file content.
    """
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > max_archive_entries:
                raise MalformedUploadError(
                    f"Uploaded DOCX has too many archive entries (limit {max_archive_entries})."
                )
            total_uncompressed = 0
            for info in infos:
                total_uncompressed += info.file_size
                if total_uncompressed > max_decompressed_bytes:
                    raise MalformedUploadError(
                        "Uploaded DOCX exceeds the maximum uncompressed size "
                        f"of {max_decompressed_bytes} bytes."
                    )
    except zipfile.BadZipFile as exc:
        # The Mime_Validator already accepted this as a DOCX, so a bytes-level
        # ZIP failure here means a malformed/truncated upload. Refuse it
        # rather than store an unreadable object.
        raise MalformedUploadError("Uploaded DOCX is not a readable archive.") from exc
